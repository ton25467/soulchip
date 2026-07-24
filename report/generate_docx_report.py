import subprocess
import datetime
import os
import re
import sys
import psutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Define all team mappings and their directory names (English & Thai)
TEAMS = {
    "structure_builder": {
        "en": "Structure Builder Team",
        "th": "ทีมพัฒนาโครงสร้างระบบหลัก"
    },
    "level_designer": {
        "en": "Level Designer Team",
        "th": "ทีมออกแบบด่านและปริศนา"
    },
    "character_designer": {
        "en": "Character Designer Team",
        "th": "ทีมออกแบบตัวละครและอนิเมชัน"
    },
    "ux_ui_designer": {
        "en": "UX/UI Designer Team",
        "th": "ทีมออกแบบส่วนต่อประสานและประสบการณ์ผู้ใช้"
    },
    "testing_team": {
        "en": "Testing Team",
        "th": "ทีมตรวจสอบคุณภาพระบบ"
    },
    "deploy_team": {
        "en": "Deploy Team",
        "th": "ทีมส่งมอบและติดตั้งระบบ"
    },
    "story_writer": {
        "en": "Story Writer Team",
        "th": "ทีมเขียนเนื้อเรื่องและออกแบบการบรรยาย"
    }
}

# Mapping of test cases to responsible teams
TEST_TEAM_MAP = {
    "TestPlayerCollision": "structure_builder",
    "TestAdversarial_CollisionAndBounds": "structure_builder",
    
    "TestLockedDoors": "level_designer",
    "TestProceduralAlgorithms": "level_designer",
    "TestStaircaseClimbing": "level_designer",
    
    "TestItemCollection": "character_designer",
    "TestSecretBossEncounter": "character_designer",
    "TestItemDrop": "character_designer",
    "TestAdversarial_PlayerInventoryWipeOnLoadLevel": "character_designer",
    
    "TestItemBoxAdjacencyAndCapacities": "ux_ui_designer",
    "TestItemBoxTransferAndSwap": "ux_ui_designer",
    "TestAdversarial_ItemBoxCapacityBoundaries": "ux_ui_designer",
    "TestAdversarial_ItemSwapLogic": "ux_ui_designer",
    "TestAdversarial_PerFloorIsolation": "ux_ui_designer",
    
    "TestCategorizedErrorReporting": "testing_team",
    
    "TestCharacterSelection": "story_writer"
}

def run_tests():
    print("Running test suite...")
    result = subprocess.run(['go', 'test', '-v', './...'], cwd='f:/soulchip', capture_output=True, text=True)
    return result.stdout, result.returncode

def parse_test_results(output):
    tests = []
    lines = output.split('\n')
    
    current_test = None
    for line in lines:
        if line.startswith("=== RUN"):
            match = re.match(r"=== RUN\s+(\w+)", line)
            if match:
                test_name = match.group(1)
                current_test = {"name": test_name, "status": "UNKNOWN", "duration": "0.00s"}
        elif line.startswith("--- PASS") or line.startswith("--- FAIL"):
            match = re.match(r"--- (PASS|FAIL):\s+(\w+)\s+\((.+)\)", line)
            if match:
                status = match.group(1)
                test_name = match.group(2)
                duration = match.group(3)
                
                test_entry = {"name": test_name, "status": status, "duration": duration}
                if not any(t['name'] == test_name for t in tests):
                    tests.append(test_entry)
                    
    alerts = []
    for line in lines:
        if "🚨 [BACKEND ALERT]" in line:
            alerts.append(line.replace("🚨 [BACKEND ALERT] ", "").strip())
            
    return tests, alerts

def get_system_metrics():
    cpu_percent = psutil.cpu_percent(interval=0.1)
    memory = psutil.virtual_memory()
    ram_total_gb = memory.total / (1024 ** 3)
    ram_used_gb = memory.used / (1024 ** 3)
    ram_percent = memory.percent
    disk = psutil.disk_usage('f:/')
    disk_total_gb = disk.total / (1024 ** 3)
    disk_used_gb = disk.used / (1024 ** 3)
    disk_percent = disk.percent
    
    return {
        "cpu_percent": cpu_percent,
        "ram_total": ram_total_gb,
        "ram_used": ram_used_gb,
        "ram_percent": ram_percent,
        "disk_total": disk_total_gb,
        "disk_used": disk_used_gb,
        "disk_percent": disk_percent
    }

def get_team_performance(team_key, tests, alerts, return_code):
    # Retrieve tests mapped to this specific team
    team_tests = [t for t in tests if TEST_TEAM_MAP.get(t['name'], "testing_team") == team_key]
    total_tests = len(team_tests)
    passed_tests = len([t for t in team_tests if t['status'] == 'PASS'])
    failed_tests = [t['name'] for t in team_tests if t['status'] == 'FAIL']
    
    # Calculate success rate
    success_rate = (passed_tests / total_tests) if total_tests > 0 else 1.0
    
    # Check for performance alert affecting UI or Physics structure
    has_performance_drop = False
    for alert in alerts:
        if "fps_drop" in alert.lower() or "not smooth" in alert.lower():
            has_performance_drop = True

    # Main build status
    is_compile_failure = (return_code != 0 and team_key == "deploy_team")
    
    # Grading Scale: A, B+, B, C+, C, D+, D, F
    if is_compile_failure:
        grade = "F"
        status_en = "F (Build / Compilation Failure)"
        status_th = "F (ระบบสร้าง Build หรือคอมไพล์โปรแกรมล้มเหลว)"
        comment_en = "Critical build compilation failure. Immediate intervention is required."
        comment_th = "ระบบคอมไพล์โปรแกรมหลักล้มเหลวโดยสิ้นเชิง ต้องได้รับการแก้ไขเป็นความสำคัญสูงสุดอันดับแรก"
    elif len(failed_tests) > 0:
        # Some tests failed - grade depends on pass rate
        if success_rate >= 0.75:
            grade = "C+"
            status_en = "C+ (Minor Failures)"
            status_th = "C+ (ผ่านการทดสอบส่วนใหญ่ แต่มีบางส่วนล้มเหลว)"
            comment_en = f"Minor failures in: {', '.join(failed_tests)}. Correcting these is Priority 1."
            comment_th = f"พบบกพร่องเล็กน้อยในจุด: {', '.join(failed_tests)} ต้องแก้ไขเป็นความสำคัญอันดับแรก (Priority 1)"
        elif success_rate >= 0.50:
            grade = "C"
            status_en = "C (Average - Moderate Failures)"
            status_th = "C (ผ่านเกณฑ์ครึ่งหนึ่ง แต่ยังมีข้อผิดพลาดหลายจุด)"
            comment_en = f"Moderate failure rate. Troubleshoot: {', '.join(failed_tests)} immediately."
            comment_th = f"พบบกพร่องปานกลางในจุด: {', '.join(failed_tests)} จำเป็นต้องวิเคราะห์และดำเนินการแก้ไขอย่างรีบด่วน"
        elif success_rate >= 0.25:
            grade = "D+"
            status_en = "D+ (Significant Failures)"
            status_th = "D+ (มีข้อผิดพลาดร้ายแรงหลายจุด ความถูกต้องต่ำกว่าเกณฑ์)"
            comment_en = f"Severe failures detected in: {', '.join(failed_tests)}. Resolution is critical."
            comment_th = f"พบข้อผิดพลาดร้ายแรงสูงในจุด: {', '.join(failed_tests)} ทีมงานต้องประเมินและทบทวนตรรกะใหม่"
        else:
            grade = "D"
            status_en = "D (Poor Performance)"
            status_th = "D (ผลงานต่ำกว่ามาตรฐานอย่างมาก การทดสอบส่วนใหญ่ไม่ผ่าน)"
            comment_en = f"Critical functionality failed: {', '.join(failed_tests)}."
            comment_th = f"โครงสร้างหลักส่วนใหญ่ไม่ผ่านการตรวจสอบ: {', '.join(failed_tests)}"
    elif has_performance_drop and team_key in ["ux_ui_designer", "structure_builder"]:
        grade = "B"
        status_en = "B (Performance Degradation)"
        status_th = "B (ระบบถูกต้องแต่มีประสิทธิภาพการทำงานล่าช้า)"
        comment_en = "FPS or TPS dropped below 60. Code optimization required."
        comment_th = "ตัวเกมรันได้ถูกต้องแต่มีเฟรมเรตตกลงต่ำกว่า 60 FPS/TPS ต้องทำการปรับปรุงลูปประสิทธิภาพ"
    elif has_performance_drop:
        grade = "B+"
        status_en = "B+ (Minor Performance Alert)"
        status_th = "B+ (การทำงานถูกต้อง สมบูรณ์แบบ แต่มีสถิติคลาดเคลื่อนเล็กน้อย)"
        comment_en = "All tests passed cleanly, but minor diagnostics alert was logged."
        comment_th = "การตรวจสอบผ่านทั้งหมด 100% แต่มีสัญญาณความล่าช้าในเครื่องเซิร์ฟเวอร์ย่อยแจ้งเตือนเข้ามา"
    else:
        grade = "A"
        status_en = "A (Excellent)"
        status_th = "A (ยอดเยี่ยมผ่านเกณฑ์ระดับสูงสุด)"
        comment_en = "Outstanding output. Zero defects, zero alerts, and optimal performance."
        comment_th = "ผลงานสมบูรณ์แบบ ไร้ข้อผิดพลาด และประสิทธิภาพการประมวลผลผ่านเกณฑ์สูงสุด"
        
    return {
        "grade": grade,
        "status_en": status_en,
        "status_th": status_th,
        "comment_en": comment_en,
        "comment_th": comment_th,
        "failed_tests": failed_tests
    }

def create_docx_report_en(stdout, return_code, tests, alerts, team_key, team_name):
    doc = Document()
    
    # 1-inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Strictly Formal EN Style: Times New Roman, 12pt, Justified
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    
    # Title (Times New Roman, Bold, 20pt)
    title = doc.add_paragraph()
    title_run = title.add_run("SOULCHIP 3D GAME DEVELOPMENT PROJECT")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(f"Official Performance & Verification Log — {team_name}")
    subtitle_run.font.name = 'Times New Roman'
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(50, 50, 50)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph() # Spacer
    
    # Metadata (Formal list)
    bangkok_tz = datetime.timezone(datetime.timedelta(hours=7))
    now_bangkok = datetime.datetime.now(bangkok_tz)
    
    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta_p.add_run("Verification Date: ").bold = True
    meta_p.add_run(now_bangkok.strftime("%d-%m-%Y %H:%M:%S (ICT - Asia/Bangkok)") + "\n")
    meta_p.add_run("Responsible Division: ").bold = True
    meta_p.add_run(f"{team_name}\n")
    meta_p.add_run("Platform Engine: ").bold = True
    meta_p.add_run("Ebitengine (Go SDK v1.26)\n")
    meta_p.add_run("Integration Verdict: ").bold = True
    
    perf = get_team_performance(team_key, tests, alerts, return_code)
    has_errors = len(perf["failed_tests"]) > 0 or (return_code != 0 and team_key == "deploy_team")
    
    verdict_run = meta_p.add_run("UNSATISFACTORY (ERRORS DETECTED)" if has_errors else "SATISFACTORY (ALL TEST CASES PASSED)")
    verdict_run.bold = True
    if has_errors:
        verdict_run.font.color.rgb = RGBColor(180, 0, 0) # Formal Dark Red
    else:
        verdict_run.font.color.rgb = RGBColor(0, 120, 0) # Formal Dark Green
        
    # Overview
    doc.add_heading("1. Executive Summary", level=1)
    summary_p = doc.add_paragraph(
        f"This official document records the daily testing results and performance diagnostics "
        f"for the '{team_name}' component. The integration and adversarial testing modules "
        f"have been verified to ensure structural compliance with design requirements."
    )
    summary_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Team Performance Evaluation Section
    doc.add_heading("2. Team Performance Evaluation", level=1)
    doc.add_paragraph("Comprehensive evaluation of the team's contribution metrics for this cycle:")
    
    p_eval = doc.add_paragraph()
    p_eval.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_eval.add_run("Evaluation Rating: ").bold = True
    g_run = p_eval.add_run(f"GRADE {perf['grade']} — {perf['status_en']}\n")
    g_run.bold = True
    
    p_eval.add_run("Evaluation Review Comments: ").bold = True
    p_eval.add_run(perf['comment_en'])
    
    # CRITICAL ACTION ITEM (PRIORITY 1) - ONLY SHOW IF ERRORS PRESENT
    if has_errors:
        doc.add_heading("🚨 3. CRITICAL DEFECT RESOLUTION MANDATE (PRIORITY 1)", level=1)
        priority_p = doc.add_paragraph()
        priority_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        priority_run = priority_p.add_run(
            "MANDATORY PROTOCOL: A quality defect has been logged during verification. "
            "Pursuant to project policy, this defect must be documented in this official report BEFORE "
            "any correction is applied. Resolving this issue constitutes the highest priority (Priority 1). "
            "All other development tasks must be halted until this defect is verified resolved."
        )
        priority_run.bold = True
        priority_run.font.color.rgb = RGBColor(180, 0, 0)
        
        doc.add_heading("Logged Defect Cause & Location:", level=2)
        if len(perf["failed_tests"]) > 0:
            for ft in perf["failed_tests"]:
                doc.add_paragraph(f"- Defect Location: test function '{ft}' inside internal/game test suites.")
                doc.add_paragraph(f"  Root Cause analysis: Logic regression in {team_name} component.")
        if return_code != 0 and team_key == "deploy_team":
            doc.add_paragraph(f"- Defect Location: Compiler toolchain execution.")
            doc.add_paragraph(f"  Root Cause analysis: Codebase build failed with compiler exit status {return_code}.")
            
    # System Resource Utilization
    doc.add_heading("4. Resource Utilization Audits", level=1)
    metrics = get_system_metrics()
    doc.add_paragraph(f"Hardware diagnostics and memory footprint audit details:")
    
    metric_table = doc.add_table(rows=1, cols=3)
    metric_table.style = 'Table Grid'
    m_hdr = metric_table.rows[0].cells
    m_hdr[0].text = 'Hardware Category'
    m_hdr[1].text = 'Allocated / Total Volume'
    m_hdr[2].text = 'Consumption Percentage'
    for cell in m_hdr:
        cell.paragraphs[0].runs[0].font.bold = True
        
    r1 = metric_table.add_row().cells
    r1[0].text = "CPU Core Utilization"
    r1[1].text = "Host Processor"
    r1[2].text = f"{metrics['cpu_percent']}%"
    
    r2 = metric_table.add_row().cells
    r2[0].text = "System RAM (Memory)"
    r2[1].text = f"{metrics['ram_used']:.2f} GB / {metrics['ram_total']:.2f} GB"
    r2[2].text = f"{metrics['ram_percent']}%"
    
    r3 = metric_table.add_row().cells
    r3[0].text = "Workspace Disk Space (F:/)"
    r3[1].text = f"{metrics['disk_used']:.2f} GB / {metrics['disk_total']:.2f} GB"
    r3[2].text = f"{metrics['disk_percent']}%"
    
    doc.add_paragraph() # Spacer

    # Test Matrix
    doc.add_heading("5. Quality Verification Matrix", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Verified Test Case'
    hdr_cells[1].text = 'Audit Status'
    hdr_cells[2].text = 'Execution Duration'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        
    for test in tests:
        row_cells = table.add_row().cells
        row_cells[0].text = test['name']
        status_cell = row_cells[1].paragraphs[0]
        status_run = status_cell.add_run(test['status'])
        status_run.bold = True
        if test['status'] == 'PASS':
            status_run.font.color.rgb = RGBColor(0, 120, 0)
        else:
            status_run.font.color.rgb = RGBColor(180, 0, 0)
        row_cells[2].text = test['duration']
        
    doc.add_paragraph() # Spacer
    
    # Raw Output
    doc.add_heading("6. System Verification Raw Log", level=1)
    raw_p = doc.add_paragraph()
    raw_run = raw_p.add_run(stdout)
    raw_run.font.name = 'Consolas'
    raw_run.font.size = Pt(8.5)
    raw_run.font.color.rgb = RGBColor(40, 40, 40)
    
    # Save EN
    team_dir = f'f:/soulchip/report/{team_key}'
    os.makedirs(team_dir, exist_ok=True)
    date_str = now_bangkok.strftime("%d-%m-%Y")
    report_path = f'{team_dir}/{date_str}.docx'
    doc.save(report_path)
    print(f"EN Report generated successfully at: {report_path}")

def create_docx_report_th(stdout, return_code, tests, alerts, team_key, team_name_th):
    doc = Document()
    
    # 1-inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Strictly Formal TH Style: TH Sarabun New, 16pt, Justified
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH Sarabun New'
    font.size = Pt(16)
    font.color.rgb = RGBColor(0, 0, 0)
    
    # Title (TH Sarabun New, Bold, 24pt)
    title = doc.add_paragraph()
    title_run = title.add_run("รายงานการทดสอบอย่างเป็นทางการ — โครงการ SOULCHIP 3D")
    title_run.font.name = 'TH Sarabun New'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(f"เอกสารบันทึกการทำงานและตรวจสอบระบบคุณภาพประจำวัน — {team_name_th}")
    subtitle_run.font.name = 'TH Sarabun New'
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(50, 50, 50)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph() # Spacer
    
    # Metadata
    bangkok_tz = datetime.timezone(datetime.timedelta(hours=7))
    now_bangkok = datetime.datetime.now(bangkok_tz)
    
    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta_p.add_run("วันที่ทดสอบระบบ: ").bold = True
    meta_p.add_run(now_bangkok.strftime("%d-%m-%Y %H:%M:%S (เวลาประเทศไทย ICT)") + "\n")
    meta_p.add_run("หน่วยงานผู้รับผิดชอบ: ").bold = True
    meta_p.add_run(f"{team_name_th}\n")
    meta_p.add_run("สภาพแวดล้อมสถาปัตยกรรม: ").bold = True
    meta_p.add_run("Ebitengine (ภาษา Go SDK 1.26)\n")
    meta_p.add_run("สรุปผลการวิเคราะห์ระบบ: ").bold = True
    
    perf = get_team_performance(team_key, tests, alerts, return_code)
    has_errors = len(perf["failed_tests"]) > 0 or (return_code != 0 and team_key == "deploy_team")
    
    verdict_run = meta_p.add_run("ไม่ผ่านเกณฑ์มาตรฐาน (พบข้อบกพร่องในระบบ)" if has_errors else "ผ่านเกณฑ์มาตรฐานการตรวจสอบความเสถียร 100%")
    verdict_run.bold = True
    if has_errors:
        verdict_run.font.color.rgb = RGBColor(180, 0, 0)
    else:
        verdict_run.font.color.rgb = RGBColor(0, 120, 0)
        
    # Overview
    doc.add_heading("1. บทสรุปภาพรวม (Executive Summary)", level=1)
    summary_p = doc.add_paragraph(
        f"เอกสารฉบับนี้เป็นรายงานอย่างเป็นทางการ เพื่อสรุปผลสัมฤทธิ์การทดสอบประจำวัน "
        f"ในส่วนรับผิดชอบของ '{team_name_th}' สำหรับระบบสลับหน้าจอสัมภาระ ด่านคฤหาสน์ 4 ชั้น "
        f"และการเคลื่อนที่ 3 มิติของตัวละครหลัก ซึ่งทั้งหมดได้รับการรันผลทดสอบเพื่อยืนยันความเสถียรเรียบร้อยแล้ว"
    )
    summary_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Team Performance Evaluation Section (TH)
    doc.add_heading("2. ผลการประเมินการทำงานของทีมงาน (Team Performance Evaluation)", level=1)
    doc.add_paragraph("การวิเคราะห์ความสามารถและประเมินผลสัมฤทธิ์ผลงานของฝ่ายพัฒนาประจำวัน:")
    
    p_eval = doc.add_paragraph()
    p_eval.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_eval.add_run("ระดับผลการประเมิน: ").bold = True
    g_run = p_eval.add_run(f"เกรด {perf['grade']} — {perf['status_th']}\n")
    g_run.bold = True
    
    p_eval.add_run("ข้อเสนอแนะของผู้ตรวจสอบระบบ: ").bold = True
    p_eval.add_run(perf['comment_th'])
    
    # CRITICAL ACTION ITEM (PRIORITY 1) - ONLY SHOW IF ERRORS PRESENT
    if has_errors:
        doc.add_heading("🚨 3. มาตรการแก้ไขข้อบกพร่องเร่งด่วนที่สุด (ความสำคัญระดับสูงสุด - Priority 1)", level=1)
        priority_p = doc.add_paragraph()
        priority_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        priority_run = priority_p.add_run(
            "ข้อบังคับสำคัญ: ตรวจพบข้อผิดพลาดจากการรันระบบตรวจสอบคุณภาพ ตามกฎระเบียบของโครงการ "
            "ทีมงานมีความจำเป็นต้องบันทึกข้อบกพร่อง ชี้แจงตำแหน่ง และสาเหตุการชำรุดลงในรายงานฉบับนี้ให้เสร็จสิ้นเสียก่อน "
            "จึงจะได้รับอนุญาตให้ทำการเริ่มแก้ไขโค้ดได้ และกระบวนการแก้ไขดังกล่าวต้องเป็นงานความสำคัญอันดับแรกสุด (Priority 1) "
            "ห้ามไม่ให้สลับไปดำเนินงานหรือฟีเจอร์อื่นๆ จนกว่าข้อผิดพลาดนี้จะถูกตรวจสอบผ่านเกณฑ์มาตรฐาน"
        )
        priority_run.bold = True
        priority_run.font.color.rgb = RGBColor(180, 0, 0)
        
        doc.add_heading("ตำแหน่งและสาเหตุข้อผิดพลาดที่ระบุ (Defect Location & Mapped Cause):", level=2)
        if len(perf["failed_tests"]) > 0:
            for ft in perf["failed_tests"]:
                doc.add_paragraph(f"- ตำแหน่งระบบ: ฟังก์ชันการรันตรวจสอบระบบ '{ft}' ในไฟล์ชุดทดสอบหลัก")
                doc.add_paragraph(f"  วิเคราะห์สาเหตุ: เกิดข้อผิดพลาดทางตรรกะระบบขอบเขตงานพัฒนาของ {team_name_th}")
        if return_code != 0 and team_key == "deploy_team":
            doc.add_paragraph(f"- ตำแหน่งระบบ: เครื่องมือประกอบโค้ดหลัก (Build Compiler)")
            doc.add_paragraph(f"  วิเคราะห์สาเหตุ: ซอร์สโค้ดภาพรวมเกิดการล้มเหลวขณะคอมไพล์โปรแกรม (รหัสความล้มเหลว {return_code})")
            
    # System Resource Utilization
    doc.add_heading("4. สถิติทรัพยากรระบบจากการตรวจสอบ (Resource Audits)", level=1)
    metrics = get_system_metrics()
    doc.add_paragraph(f"สรุปการใช้ทรัพยากรของเครื่องทดสอบหลัก ณ เวลาการประมวลผล:")
    
    metric_table = doc.add_table(rows=1, cols=3)
    metric_table.style = 'Table Grid'
    m_hdr = metric_table.rows[0].cells
    m_hdr[0].text = 'ส่วนประกอบของระบบ'
    m_hdr[1].text = 'ปริมาณการใช้ปัจจุบัน / ความจุรวมสูงสุด'
    m_hdr[2].text = 'สัดส่วนเปอร์เซ็นต์'
    for cell in m_hdr:
        cell.paragraphs[0].runs[0].font.bold = True
        
    r1 = metric_table.add_row().cells
    r1[0].text = "อัตราประมวลผลของหน่วยประมวลผลกลาง"
    r1[1].text = "หน่วยประมวลผลหลัก (CPU)"
    r1[2].text = f"{metrics['cpu_percent']}%"
    
    r2 = metric_table.add_row().cells
    r2[0].text = "อัตราการจองหน่วยความจำชั่วคราว"
    r2[1].text = f"{metrics['ram_used']:.2f} GB / {metrics['ram_total']:.2f} GB"
    r2[2].text = f"{metrics['ram_percent']}%"
    
    r3 = metric_table.add_row().cells
    r3[0].text = "อัตราพื้นที่จัดเก็บฮาร์ดดิสก์โครงการ"
    r3[1].text = f"{metrics['disk_used']:.2f} GB / {metrics['disk_total']:.2f} GB"
    r3[2].text = f"{metrics['disk_percent']}%"
    
    doc.add_paragraph() # Spacer

    # Test Matrix
    doc.add_heading("5. รายละเอียดตารางบันทึกการตรวจสอบคุณภาพ", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'หัวข้อตรวจสอบคุณภาพ'
    hdr_cells[1].text = 'ผลลัพธ์การประเมิน'
    hdr_cells[2].text = 'ระยะเวลาประมวลผล'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        
    for test in tests:
        row_cells = table.add_row().cells
        row_cells[0].text = test['name']
        status_cell = row_cells[1].paragraphs[0]
        status_run = status_cell.add_run("ผ่านเกณฑ์" if test['status'] == 'PASS' else "ชำรุด")
        status_run.bold = True
        if test['status'] == 'PASS':
            status_run.font.color.rgb = RGBColor(0, 120, 0)
        else:
            status_run.font.color.rgb = RGBColor(180, 0, 0)
        row_cells[2].text = test['duration']
        
    doc.add_paragraph() # Spacer
    
    # Raw Output
    doc.add_heading("6. บันทึกผลลัพธ์การรันโปรแกรมระบบแบบดิบ (Raw Log)", level=1)
    raw_p = doc.add_paragraph()
    raw_run = raw_p.add_run(stdout)
    raw_run.font.name = 'Consolas'
    raw_run.font.size = Pt(8.5)
    raw_run.font.color.rgb = RGBColor(40, 40, 40)
    
    # Save TH
    team_dir = f'f:/soulchip/report/{team_key}'
    os.makedirs(team_dir, exist_ok=True)
    date_str = now_bangkok.strftime("%d-%m-%Y")
    report_path = f'{team_dir}/รายงานฉบับ1_{date_str}.docx'
    doc.save(report_path)
    print(f"TH Report generated successfully for {team_key}")

if __name__ == '__main__':
    stdout, return_code = run_tests()
    tests, alerts = parse_test_results(stdout)
    
    if len(sys.argv) > 1:
        team_key = sys.argv[1]
        if team_key in TEAMS:
            create_docx_report_en(stdout, return_code, tests, alerts, team_key, TEAMS[team_key]["en"])
            create_docx_report_th(stdout, return_code, tests, alerts, team_key, TEAMS[team_key]["th"])
        else:
            print(f"Unknown team key: {team_key}. Available teams: {list(TEAMS.keys())}")
            sys.exit(1)
    else:
        print("No specific team specified. Generating reports (EN & TH) for all teams...")
        for team_key, team_info in TEAMS.items():
            create_docx_report_en(stdout, return_code, tests, alerts, team_key, team_info["en"])
            create_docx_report_th(stdout, return_code, tests, alerts, team_key, team_info["th"])
